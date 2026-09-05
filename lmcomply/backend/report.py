import json
from datetime import datetime
from fpdf import FPDF


def _safe(value):
    if value is None:
        return "Not detected"
    if isinstance(value, dict):
        return _safe(value.get("value"))
    if isinstance(value, (tuple, list)):
        return " ".join(map(str, value))
    return str(value).replace("₹", "Rs.")


def _data(scan):
    return json.loads(scan["extracted"] or "{}")


def build_pdf(scan, product, violations, out_path):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page(); pdf.set_title("Legal Metrology Compliance Report")
    pdf.set_font("helvetica", "B", 18); pdf.cell(0, 10, "Legal Metrology Compliance Report", ln=True, align="C")
    pdf.set_font("helvetica", "", 9); pdf.cell(0, 6, "AI-assisted label screening | Inspector verification required", ln=True, align="C"); pdf.ln(5)
    pdf.set_font("helvetica", "B", 11); pdf.cell(0, 7, "Inspection Summary", ln=True)
    pdf.set_font("helvetica", "", 10)
    pdf.cell(0, 6, f"Product: {_safe(product['name'])}", ln=True); pdf.cell(0, 6, f"Category: {_safe(product['category'])}", ln=True)
    pdf.cell(0, 6, f"Scan ID: {scan['id']}   Status: {scan['status']}", ln=True)
    try: created = datetime.fromtimestamp(float(scan['created_at'])).strftime('%d-%m-%Y %H:%M')
    except Exception: created = _safe(scan['created_at'])
    pdf.cell(0, 6, f"Inspection time: {created}", ln=True); pdf.ln(3)
    ex = _data(scan); pdf.set_font("helvetica", "B", 11); pdf.cell(0, 7, "Declaration Detection", ln=True); pdf.set_font("helvetica", "", 9)
    for label, key in [("Product name","product_name"),("MRP","mrp"),("Net quantity","net_qty"),("Manufacture/Packing date","date"),("Best before","best_before"),("Manufacturer / packer","manufacturer"),("Importer","importer"),("Consumer care","consumer_care"),("Country of origin","origin"),("OCR confidence","ocr_confidence"),("Estimated MRP text height","font_mm")]:
        pdf.multi_cell(0, 5, f"{label}: {_safe(ex.get(key))}")
    pdf.ln(3); pdf.set_font("helvetica", "B", 11); pdf.cell(0, 7, f"Findings ({len(violations)})", ln=True); pdf.set_font("helvetica", "", 9)
    if not violations: pdf.multi_cell(0, 5, "No critical or major issues were flagged by the configured screening rules.")
    for v in violations: pdf.multi_cell(0, 5, f"[{v['severity']}] {v['rule_id']}: {_safe(v['message'])}")
    pdf.ln(4); pdf.set_font("helvetica", "B", 10); pdf.cell(0, 6, "Inspector action", ln=True); pdf.set_font("helvetica", "", 8)
    pdf.multi_cell(0, 4, "Verify declaration placement, applicable type size, readability, actual quantity, batch/date details and category-specific requirements against the package and current Legal Metrology requirements. Automated screening is not legal adjudication.")
    pdf.output(out_path)


def build_docx(scan, product, violations, out_path):
    from docx import Document
    doc = Document(); doc.add_heading("Legal Metrology Compliance Report", 0)
    doc.add_paragraph("AI-assisted label screening | Inspector verification required")
    doc.add_heading("Inspection Summary", level=1)
    doc.add_paragraph(f"Product: {_safe(product['name'])}\nCategory: {_safe(product['category'])}\nScan ID: {scan['id']}\nStatus: {scan['status']}")
    doc.add_heading("Declaration Detection", level=1)
    ex = _data(scan)
    table = doc.add_table(rows=1, cols=2); table.style = "Table Grid"
    table.rows[0].cells[0].text = "Declaration"; table.rows[0].cells[1].text = "Detected value"
    for label, key in [("Product name","product_name"),("MRP","mrp"),("Net quantity","net_qty"),("Manufacture/Packing date","date"),("Best before","best_before"),("Manufacturer / packer","manufacturer"),("Importer","importer"),("Consumer care","consumer_care"),("Country of origin","origin"),("OCR confidence","ocr_confidence"),("Estimated MRP text height","font_mm")]:
        cells = table.add_row().cells; cells[0].text = label; cells[1].text = _safe(ex.get(key))
    doc.add_heading(f"Findings ({len(violations)})", level=1)
    if not violations: doc.add_paragraph("No critical or major issues were flagged by the configured screening rules.")
    for v in violations: doc.add_paragraph(f"[{v['severity']}] {v['rule_id']}: {_safe(v['message'])}", style="List Bullet")
    doc.add_heading("Inspector action", level=1)
    doc.add_paragraph("Verify declaration placement, applicable type size, readability, actual quantity, batch/date details and category-specific requirements against the package and current Legal Metrology requirements. Automated screening is not legal adjudication.")
    doc.save(out_path)
